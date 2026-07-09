#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
export_docx.py — Exegete 주해 Markdown → Word(.docx) 변환기

주해 결과(output/*.md)를 서식이 살아있는 워드 문서로 변환한다.
순수 파이썬(python-docx)만 사용하므로 외부 렌더러(LibreOffice 등)가 필요 없고,
Windows·macOS·Linux 어디서나 동일하게 동작한다. 만든 .docx는 MS Word뿐 아니라
한컴 한글에서도 그대로 열려 편집할 수 있다.

지원하는 마크다운 요소: 제목(#/##/###), 볼드/이탤릭/인라인코드,
인용블록(>), 표(|), 코드블록(```) = 구조 도표, 순서/비순서 리스트, 구분선(---).

사용법:
    python src/export_docx.py output/엡2_8-9_주해.md
    python src/export_docx.py output/엡2_8-9_주해.md -o C:/some/where/out.docx
    python src/export_docx.py --all          # output/ 안의 모든 .md 를 .docx 로
"""
import sys
import os
import re
import glob
import argparse

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "python-docx 가 필요합니다.  설치:  pip install python-docx\n"
    )
    sys.exit(1)

# ── 서식 상수 ───────────────────────────────────────────────
BODY_FONT = "맑은 고딕"      # 한글 본문
MONO_FONT = "Consolas"       # 코드블록·구조 도표(고정폭 → 정렬 보존)
HEB_FONT = "Times New Roman"  # 히브리어 등 complex script(오른쪽→왼쪽) 전용 폰트

C_TITLE = RGBColor(0x1F, 0x38, 0x64)   # 문서 제목 남색
C_H1    = RGBColor(0x2E, 0x4B, 0x7A)   # 대제목
C_H2    = RGBColor(0x3B, 0x5A, 0x8A)   # 소제목
C_CODE  = RGBColor(0x33, 0x33, 0x33)   # 코드 텍스트
C_ICODE = RGBColor(0xC0, 0x30, 0x30)   # 인라인 코드(붉은 갈색)
C_QUOTE = RGBColor(0x33, 0x3A, 0x45)   # 인용 텍스트

HEX_TITLE = "1F3864"
HEX_H1    = "2E4B7A"
SH_QUOTE  = "F2F4F8"   # 인용 배경
SH_CODE   = "F4F3EE"   # 코드블록 배경
SH_THEAD  = "DCE6F1"   # 표 머리행 음영
BD_QUOTE  = "4A86C4"   # 인용 좌측 테두리
BD_CODE   = "CCCCCC"   # 코드블록 테두리
HR_COLOR  = "BBBBBB"   # 구분선


# ── 저수준 헬퍼(폰트·음영·테두리) ─────────────────────────────
def _apply_run_font(run, name, cs_name=None):
    """run 에 ascii/hAnsi/eastAsia/cs 폰트를 지정. cs_name 으로 히브리어 등
    complex script 전용 폰트를 따로 줄 수 있다(한글·히브리어 동시 적용의 핵심)."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), cs_name or name)


def _set_cs_size(run, pt):
    """complex script(히브리어 등)용 글자 크기(szCs)를 본문과 맞춘다.
    지정하지 않으면 히브리어가 기본 크기로 작게 렌더된다."""
    rPr = run._element.get_or_add_rPr()
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(int(round(pt * 2))))


def _apply_style_font(style, name, cs_name=None, size_pt=None):
    style.font.name = name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), cs_name or name)
    if size_pt is not None:
        for tag in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn("w:val"), str(int(round(size_pt * 2))))


def _shade_paragraph(paragraph, hex_fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _para_border(paragraph, edges, hex_color, sz="8", space="4"):
    """edges: ('left',) 또는 ('top','left','bottom','right') 등."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for edge in edges:
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), space)
        el.set(qn("w:color"), hex_color)
        pBdr.append(el)


# ── 인라인(**볼드**, *이탤릭*, `코드`) ────────────────────────
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*\n]+?\*)")


def _add_run(paragraph, text, *, bold=False, italic=False, mono=False,
             size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if mono:
        _apply_run_font(run, MONO_FONT, cs_name=MONO_FONT)
    else:
        _apply_run_font(run, BODY_FONT, cs_name=HEB_FONT)
    if size is not None:
        run.font.size = Pt(size)
        _set_cs_size(run, size)
    if color is not None:
        run.font.color.rgb = color
    return run


def _add_inline(paragraph, text, *, base_bold=False, size=None, color=None):
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _add_run(paragraph, text[pos:m.start()],
                     bold=base_bold, size=size, color=color)
        s = m.group()
        if s.startswith("**"):
            _add_run(paragraph, s[2:-2], bold=True, size=size, color=color)
        elif s.startswith("`"):
            _add_run(paragraph, s[1:-1], mono=True, size=size, color=C_ICODE)
        else:  # *이탤릭*
            _add_run(paragraph, s[1:-1], italic=True, bold=base_bold,
                     size=size, color=color)
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], bold=base_bold, size=size, color=color)


# ── 블록 요소 ────────────────────────────────────────────────
def _add_heading(doc, text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level <= 1:
        pf.space_before = Pt(2)
        pf.space_after = Pt(10)
        _add_inline(p, text, size=19, color=C_TITLE, base_bold=True)
        _para_border(p, ("bottom",), HEX_TITLE, sz="18", space="4")
    elif level == 2:
        pf.space_before = Pt(14)
        pf.space_after = Pt(6)
        pf.keep_with_next = True
        _add_inline(p, text, size=14, color=C_H1, base_bold=True)
        _para_border(p, ("bottom",), HEX_H1, sz="6", space="3")
    else:
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        pf.keep_with_next = True
        _add_inline(p, text, size=12, color=C_H2, base_bold=True)


def _add_quote(doc, qlines):
    # 앞뒤 빈 줄 정리
    while qlines and qlines[0].strip() == "":
        qlines.pop(0)
    while qlines and qlines[-1].strip() == "":
        qlines.pop()
    for j, q in enumerate(qlines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.22)
        pf.space_after = Pt(2 if j < len(qlines) - 1 else 8)
        pf.space_before = Pt(8 if j == 0 else 0)
        _shade_paragraph(p, SH_QUOTE)
        _para_border(p, ("left",), BD_QUOTE, sz="18", space="8")
        _add_inline(p, q if q.strip() else " ", size=10, color=C_QUOTE)


def _add_code_block(doc, code_lines):
    while code_lines and code_lines[-1].strip() == "":
        code_lines.pop()
    if not code_lines:
        return
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    _shade_paragraph(p, SH_CODE)
    _para_border(p, ("top", "left", "bottom", "right"), BD_CODE, sz="6", space="6")
    run = p.add_run()
    _apply_run_font(run, MONO_FONT)
    run.font.size = Pt(9.5)
    run.font.color.rgb = C_CODE
    for k, cl in enumerate(code_lines):
        if k > 0:
            run.add_break()
        run.add_text(cl)


_SEP_RE = re.compile(r"^\|[\s:\-|]+\|$")


def _add_table(doc, tlines):
    rows = []
    for tl in tlines:
        if _SEP_RE.match(tl):
            continue
        cells = [c.strip() for c in tl.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = "Table Grid"
    table.autofit = True
    for ri, r in enumerate(rows):
        for ci in range(ncol):
            cell = table.cell(ri, ci)
            p = cell.paragraphs[0]
            p.text = ""
            txt = r[ci] if ci < len(r) else ""
            _add_inline(p, txt, base_bold=(ri == 0), size=9.5)
            if ri == 0:
                _shade_cell(cell, SH_THEAD)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_list_item(doc, text, ordered):
    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
        _add_run(p, "• ")
    _add_inline(p, text)


def _add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _add_inline(p, text)


# ── 메인 변환 ────────────────────────────────────────────────
def convert(md_path, out_path):
    with open(md_path, encoding="utf-8") as f:
        md = f.read()
    lines = md.split("\n")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    _apply_style_font(normal, BODY_FONT, cs_name=HEB_FONT, size_pt=10.5)
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    i, n = 0, len(lines)
    while i < n:
        stripped = lines[i].strip()

        # 코드블록
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # 닫는 ```
            _add_code_block(doc, buf)
            continue

        # 표
        if stripped.startswith("|") and stripped.endswith("|") and len(stripped) > 1:
            buf = []
            while i < n and lines[i].strip().startswith("|"):
                buf.append(lines[i].strip())
                i += 1
            _add_table(doc, buf)
            continue

        # 빈 줄
        if stripped == "":
            i += 1
            continue

        # 구분선
        if re.fullmatch(r"-{3,}", stripped) or re.fullmatch(r"\*{3,}", stripped) \
                or re.fullmatch(r"_{3,}", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _para_border(p, ("bottom",), HR_COLOR, sz="6", space="1")
            i += 1
            continue

        # 제목
        m = re.match(r"(#{1,6})\s+(.*)", stripped)
        if m:
            _add_heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue

        # 인용(연속 > 모으기)
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip()[1:].lstrip())
                i += 1
            _add_quote(doc, buf)
            continue

        # 비순서 리스트
        if re.match(r"[-*+]\s+", stripped):
            _add_list_item(doc, stripped[2:].strip(), ordered=False)
            i += 1
            continue

        # 순서 리스트
        m = re.match(r"\d+\.\s+(.*)", stripped)
        if m:
            _add_list_item(doc, m.group(1).strip(), ordered=True)
            i += 1
            continue

        # 일반 문단
        _add_paragraph(doc, stripped)
        i += 1

    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser(
        description="Exegete 주해 Markdown → Word(.docx) 변환")
    ap.add_argument("md", nargs="?", help="변환할 .md 파일 경로")
    ap.add_argument("-o", "--out", help="출력 .docx 경로(단일 변환 시)")
    ap.add_argument("--all", action="store_true",
                    help="output/ 안의 모든 .md 를 변환")
    args = ap.parse_args()

    repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    if args.all:
        targets = sorted(glob.glob(os.path.join(repo_root, "output", "*.md")))
        if not targets:
            print("output/ 에 변환할 .md 가 없습니다.")
            return
    elif args.md:
        targets = [args.md]
    else:
        ap.print_help()
        return

    ok = 0
    for md_path in targets:
        if not os.path.isfile(md_path):
            print(f"✗ 파일 없음: {md_path}")
            continue
        if args.out and not args.all:
            out = args.out
        else:
            out = os.path.splitext(md_path)[0] + ".docx"
        try:
            convert(md_path, out)
            print(f"✓ {os.path.basename(md_path)} → {os.path.basename(out)}")
            ok += 1
        except Exception as e:
            print(f"✗ {os.path.basename(md_path)} 변환 실패: {e}")
    print(f"\n완료: {ok}/{len(targets)} 개 변환")


if __name__ == "__main__":
    main()
